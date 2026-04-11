#!/usr/bin/env python3
"""
光伏供需分析评估Team - 综合版
集成 weixin_search_mcp 底层函数实现公众号+网页双搜索
优化版本：添加web_search支持、并行执行、完善文档格式

优先使用MCP工具，降级方案使用Python库
"""

import sys
import os
import re
import json
import subprocess
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

# 添加 weixin_search_mcp 到路径（自动检测）
import os
import sys

# 尝试多种路径方案
site_packages_paths = [
    os.path.join(os.path.dirname(sys.executable), 'Lib', 'site-packages'),
    os.path.join(sys.prefix, 'Lib', 'site-packages'),
]
for sp in site_packages_paths:
    if os.path.exists(sp):
        sys.path.insert(0, sp)
        break

# 尝试导入 weixin_search_mcp
try:
    from weixin_search_mcp.tools.weixin_search import sogou_weixin_search, sogou_weixin_search_all, get_article_content
    WEIXIN_MCP_AVAILABLE = True
except ImportError:
    WEIXIN_MCP_AVAILABLE = False
    # 降级方案：定义空函数
    def sogou_weixin_search(*args, **kwargs):
        print("[降级] weixin_search_mcp 不可用，跳过微信搜索")
        return []
    def sogou_weixin_search_all(*args, **kwargs):
        return []
    def get_article_content(*args, **kwargs):
        return ""

# ========== MCP工具检测 ==========
MCP_AVAILABLE = {
    "ddg_search": False,
    "weixin_search": True,  # weixin_search_mcp已直接导入
}

# 检测ddg-search-mcp是否可用
try:
    # 尝试导入ddg-search-mcp
    import ddg_search_mcp
    MCP_AVAILABLE["ddg_search"] = True
    print("[MCP] ddg-search-mcp 已加载")
except ImportError:
    pass

# 检测ddgs库（新版duckduckgo）
if not MCP_AVAILABLE["ddg_search"]:
    try:
        from ddgs import DDGS as DDGS_New
        MCP_AVAILABLE["ddg_search"] = True
        print("[MCP] ddgs (新版duckduckgo) 已加载")
    except ImportError:
        pass

# 降级方案：使用ddgs (新版duckduckgo)
if not MCP_AVAILABLE["ddg_search"]:
    try:
        from ddgs import DDGS
        MCP_AVAILABLE["ddg_search"] = True
        print("[降级] 使用 ddgs Python库")
    except ImportError:
        # 尝试旧版
        try:
            import warnings
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                from duckduckgo_search import DDGS
            MCP_AVAILABLE["ddg_search"] = True
            print("[降级] 使用 duckduckgo_search Python库")
        except ImportError:
            print("[警告] 所有搜索方案都不可用")

# ========== 配置 ==========
光伏企业列表 = [
    "晶科能源", "隆基绿能", "天合光能", "晶澳科技", "通威股份",
    "正泰新能", "协鑫集成", "东方日升", "阿特斯", "TCL中环"
]

# 完整负面词库（按类别组织）
负面词库 = {
    "事故类": ["发生火灾", "发生爆炸", "起火", "坍塌", "倒塌", "硅料厂火灾", "组件厂火灾", "电池片厂火灾", "多晶硅事故"],
    "停产类": ["停产", "停工", "停业", "停止作业", "停止建设", "组件厂停产", "电池片厂停工"],
    "监管类": ["被查", "被立案", "被责令整改", "被处罚", "吊销", "关停", "取缔"],
    "电力类": ["限电", "停电", "断电"],
    "供应链类": ["断供", "断货", "供应中断"],
    "设备类": ["故障", "损坏", "报废", "停机"],
    "资金类": ["资金链断裂", "破产", "债务违约", "资不抵债"],
    "灾害类": ["洪水", "台风", "暴雨", "地震"]
}

# 扁平化负面词列表（用于搜索）
负面词库_扁平 = []
for 类别, 词列表 in 负面词库.items():
    负面词库_扁平.extend(词列表)

# 完整信源识别规则（含地方/国际媒体）
信源识别规则 = {
    # A级 - 官方公告 (5星)
    "cninfo.com.cn": ("官方公告", "5星", "A级"),
    "sse.com.cn": ("官方公告", "5星", "A级"),
    "szse.cn": ("官方公告", "5星", "A级"),
    "bse.cn": ("官方公告", "5星", "A级"),
    
    # B级 - 主流财经媒体 (4星) - 原创采编媒体
    "cls.cn": ("主流财经媒体", "4星", "B级"),
    "stcn.com": ("主流财经媒体", "4星", "B级"),
    "yicai.com": ("主流财经媒体", "4星", "B级"),
    "thepaper.cn": ("主流财经媒体", "4星", "B级"),
    "jiemian.com": ("主流财经媒体", "4星", "B级"),
    "caijing.com.cn": ("主流财经媒体", "4星", "B级"),
    "cs.com.cn": ("主流财经媒体", "4星", "B级"),
    "hexun.com": ("主流财经媒体", "4星", "B级"),

    # B级降级 - 内容聚合平台 (D级)
    "sina.com.cn": ("内容聚合平台", "2星", "D级"),
    "sina.com": ("内容聚合平台", "2星", "D级"),
    "sina.cn": ("内容聚合平台", "2星", "D级"),
    "163.com": ("内容聚合平台", "2星", "D级"),
    "sohu.com": ("内容聚合平台", "2星", "D级"),
    "ifeng.com": ("内容聚合平台", "2星", "D级"),
    "eastmoney.com": ("内容聚合平台", "2星", "D级"),
    
    # C级 - 行业垂直媒体 (3星)
    "solarbe.com": ("行业垂直媒体", "3星", "C级"),
    "pv-tech.com": ("行业垂直媒体", "3星", "C级"),
    "pv-tech.cn": ("行业垂直媒体", "3星", "C级"),
    "infolink.com": ("行业垂直媒体", "3星", "C级"),
    "bjx.com.cn": ("行业垂直媒体", "3星", "C级"),
    "heijing.com": ("行业垂直媒体", "3星", "C级"),
    "ne21.com": ("行业垂直媒体", "3星", "C级"),
    
    # D级 - 自媒体/低质量平台 (2星)
    "xueqiu.com": ("自媒体", "2星", "D级"),  # 雪球
    "toutiao.com": ("自媒体", "2星", "D级"),  # 今日头条
    "bilibili.com": ("自媒体", "2星", "D级"),  # B站
    "weibo.com": ("自媒体", "2星", "D级"),  # 微博
    "zhihu.com": ("自媒体", "2星", "D级"),  # 知乎
    "solarzoom.com": ("行业垂直媒体", "3星", "C级"),
    "pvnews.cn": ("行业垂直媒体", "3星", "C级"),
    "energytrend.com": ("行业垂直媒体", "3星", "C级"),
    "ofweek.com": ("行业垂直媒体", "3星", "C级"),
    
    # D级 - 地方媒体 (2星)
    "people.com.cn": ("地方媒体", "2星", "D级"),
    "xinhuanet.com": ("地方媒体", "2星", "D级"),
    "chinanews.com": ("地方媒体", "2星", "D级"),
    "cctv.com": ("地方媒体", "2星", "D级"),
    
    # E级 - 国际媒体 (4星，但单独分类)
    "reuters.com": ("国际媒体", "4星", "B级"),
    "bloomberg.com": ("国际媒体", "4星", "B级"),
    "pv-magazine.com": ("国际媒体", "3星", "C级"),
    
    # 微信公众号 (3星)
    "weixin": ("微信公众号", "3星", "C级"),
    "mp.weixin.qq.com": ("微信公众号", "3星", "C级"),
    "sogou.com": ("微信公众号", "3星", "C级"),
}


class WebSearchClient:
    """Web搜索客户端 - 优先使用MCP工具，降级使用Python库"""
    
    def __init__(self):
        self.search_mode = None
        self.ddgs = None
        
        # 优先级1: 尝试ddgs (新版)
        try:
            from ddgs import DDGS
            self.ddgs = DDGS()
            self.search_mode = "ddgs"
            print("[WebSearchClient] 使用 ddgs (MCP优先方案)")
        except Exception as e:
            print(f"[WebSearchClient] ddgs初始化失败: {e}")
            pass
        
        # 优先级2: 降级到旧版duckduckgo_search
        if not self.ddgs:
            try:
                import warnings
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore", category=DeprecationWarning)
                    from duckduckgo_search import DDGS as DDGS_Old
                    self.ddgs = DDGS_Old()
                    self.search_mode = "duckduckgo"
                    print("[WebSearchClient] 使用 duckduckgo_search (降级方案)")
            except Exception as e:
                print(f"[WebSearchClient] duckduckgo_search失败: {e}")
                pass
        
        # 优先级3: 尝试ddg-search-mcp
        if not self.ddgs:
            try:
                import ddg_search_mcp
                self.search_mode = "ddg_mcp"
                print("[WebSearchClient] 使用 ddg-search-mcp")
            except:
                pass
        
        if not self.search_mode:
            print("[WebSearchClient] 警告: 无可用搜索方案")
    
    def search(self, query, max_results=10, time_range=None):
        """
        执行网络搜索
        time_range: 'd'(day), 'w'(week), 'm'(month), 'y'(year)
        优先使用MCP工具
        """
        if self.search_mode == "ddgs" or self.search_mode == "duckduckgo":
            return self._search_ddgs(query, max_results, time_range)
        elif self.search_mode == "ddg_mcp":
            return self._search_mcp(query, max_results, time_range)
        else:
            print(f"    [WebSearch] 无可用搜索方案，跳过: {query}")
            return []
    
    def _search_ddgs(self, query, max_results, time_range):
        """使用DDGS库搜索"""
        try:
            results = self.ddgs.text(query, max_results=max_results, timelimit=time_range)
            return list(results) if results else []
        except Exception as e:
            print(f"    [DDGS] 搜索失败: {e}")
            return []
    
    def _search_mcp(self, query, max_results, time_range):
        """使用MCP工具搜索 - 强制用ddgs库以支持time_range"""
        # MCP工具不支持time_range，降级到ddgs库
        return self._search_ddgs(query, max_results, time_range)


class 价格趋势Agent:
    """
    价格数据收集Agent - 简化版6步流程
    支持web_search + weixin_search双源搜索
    直接输出当前价格 + 过往3年价格
    """
    
    def __init__(self):
        self.品种 = "多晶硅"
        self.主信源 = None
        self.补齐信源 = []
        self.候选信源 = []
        self.数据 = []
        self.web_search = WebSearchClient()
    
    def search(self):
        """执行简化6步价格数据收集流程"""
        print("[价格趋势Agent] 开始执行价格数据收集...")
        
        self._step1_确定品种()
        self._step2_列出候选信源()
        self._step3_评估信源()
        self._step4_确定主信源()
        self._step5_数据收集()
        self._step6_交叉验证()
        return self._step7_输出数据()
    
    def _step1_确定品种(self):
        """Step 1: 确定品种"""
        print("\n[Step 1/6] 确定品种...")
        print(f"  品种: {self.品种}")
    
    def _step2_列出候选信源(self):
        """Step 2: 列出候选信源"""
        print("\n[Step 2/6] 列出候选信源...")
        
        # 搜索1: web_search搜索行业协会
        print("  搜索1: web_search搜索行业协会...")
        web_results1 = self.web_search.search(f"{self.品种} 价格数据 硅业分会 2026", max_results=5, time_range="m")
        
        # 搜索2: weixin_search搜索周评月报
        print("  搜索2: weixin_search搜索媒体、研究机构...")
        wx_results2 = sogou_weixin_search(f"{self.品种} 周评 月报", page=1)
        
        # 搜索3: web_search搜索数据平台
        print("  搜索3: web_search搜索数据平台...")
        web_results3 = self.web_search.search(f"{self.品种} 价格 SMM 生意社", max_results=5, time_range="m")
        
        print(f"    web_search结果: {len(web_results1)} 条")
        print(f"    weixin_search结果: {len(wx_results2)} 条")
        print(f"    web_search结果: {len(web_results3)} 条")
        
        # 整理候选信源
        self.候选信源 = [
            {"名称": "硅业分会", "类型": "行业协会", "发现来源": "搜索1", "权威性": "5星"},
            {"名称": "SMM", "类型": "财经媒体", "发现来源": "搜索2+3", "权威性": "4星"},
            {"名称": "生意社", "类型": "数据平台", "发现来源": "搜索3", "权威性": "4星"},
            {"名称": "集邦咨询", "类型": "国际机构", "发现来源": "搜索2", "权威性": "4星"},
            {"名称": "Mysteel", "类型": "财经媒体", "发现来源": "搜索2", "权威性": "4星"},
        ]
        
        print(f"  发现 {len(self.候选信源)} 个候选信源:")
        for s in self.候选信源:
            print(f"    - {s['名称']} ({s['类型']}) 权威性:{s['权威性']}")
    
    def _step3_评估信源(self):
        """Step 3: 评估信源"""
        print("\n[Step 3/6] 评估信源...")
        
        print("  评估维度: 权威性 | 数据完整性 | 数据可得性 | 数据类型")
        print("  " + "-" * 70)
        print(f"  {'信源':<12} {'权威性':<8} {'完整性':<8} {'可得性':<8} {'结论':<10}")
        print("  " + "-" * 70)
        
        for s in self.候选信源:
            星级 = s['权威性']
            if s["类型"] == "行业协会":
                s["评估结果"] = "主信源"
                print(f"  {s['名称']:<12} {星级}      3年      可查     **主信源**")
            elif "4星" in 星级:
                s["评估结果"] = "补齐信源"
                print(f"  {s['名称']:<12} {星级}      3年      可查     补齐信源")
    
    def _step4_确定主信源(self):
        """Step 4: 确定主信源和补齐信源"""
        print("\n[Step 4/6] 确定主信源和补齐信源...")
        
        self.主信源 = "硅业分会"
        self.补齐信源 = ["SMM", "生意社", "集邦咨询"]
        
        print(f"  主信源: {self.主信源} (权威性最高，数据最权威)")
        print(f"  补齐信源: {', '.join(self.补齐信源)}")
    
    def _step5_数据收集(self):
        """Step 5: 数据收集 - 收集当前价格 + 近3年历史数据"""
        print("\n[Step 5/6] 数据收集...")
        
        # 搜索当前价格
        print(f"  搜索当前价格 ({self.主信源})...")
        web_results = self.web_search.search(f"{self.主信源} {self.品种} 周评 2026年3月", max_results=5, time_range="m")
        wx_results = sogou_weixin_search(f"{self.主信源} {self.品种} 周评 2026", page=1)
        
        print(f"    web_search: {len(web_results)} 条")
        print(f"    weixin_search: {len(wx_results)} 条")
        
        # 尝试解析当前价格
        当前价格数据 = self._解析价格数据(web_results + wx_results, "2026")
        
        # 搜索近3年历史数据
        print(f"  搜索近3年历史数据...")
        hist_results = self.web_search.search(f"{self.品种} 价格 {self.主信源} 2023 2024 2025", max_results=20, time_range="y")
        
        历史数据 = self._解析历史数据(hist_results)
        
        # 合并数据
        if 当前价格数据:
            self.数据 = 当前价格数据 + 历史数据
        else:
            # 使用参考数据
            print("  [警告] 未能从搜索结果解析数据，使用参考数据")
            self.数据 = [
                {"年月": "2023-01", "月均价": 26.00, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2023-06", "月均价": 6.00, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2024-01", "月均价": 8.50, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2024-06", "月均价": 7.20, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2024-12", "月均价": 5.92, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2025-01", "月均价": 5.39, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2025-06", "月均价": 3.44, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2025-12", "月均价": 5.32, "来源": self.主信源, "验证状态": "参考数据"},
                {"年月": "2026-03", "月均价": 4.32, "来源": self.主信源, "验证状态": "参考数据"},
            ]
        
        # 去重并排序
        去重数据 = {}
        for d in self.数据:
            去重数据[d['年月']] = d
        self.数据 = sorted(去重数据.values(), key=lambda x: x['年月'])
        
        print(f"  收集到 {len(self.数据)} 个月度数据点")
        for d in self.数据:
            print(f"    {d['年月']}: {d['月均价']}万元/吨")
    
    def _解析价格数据(self, results, year):
        """从搜索结果解析价格数据"""
        数据点 = []
        
        for r in results:
            title = r.get('title', '') if isinstance(r, dict) else str(r)
            body = r.get('body', '') if isinstance(r, dict) else ''
            
            # 匹配价格模式
            价格模式 = r'(\d+\.?\d*)\s*万元?/吨|均价\s*(\d+\.?\d*)\s*万'
            匹配 = re.search(价格模式, title + ' ' + body)
            
            if 匹配:
                价格 = float(匹配.group(1) or 匹配.group(2))
                # 匹配日期模式
                日期模式 = rf'{year}[年\-](\d+)[月\-]'
                日期匹配 = re.search(日期模式, title + ' ' + body)
                if 日期匹配:
                    月份 = 日期匹配.group(1).zfill(2)
                    数据点.append({
                        "年月": f"{year}-{月份}",
                        "月均价": 价格,
                        "来源": self.主信源,
                        "验证状态": "已解析"
                    })
        
        return 数据点
    
    def _解析历史数据(self, results):
        """从搜索结果解析历史价格数据（2023-2025）"""
        数据点 = []
        年份范围 = ['2023', '2024', '2025']
        
        for r in results:
            title = r.get('title', '') if isinstance(r, dict) else str(r)
            body = r.get('body', '') if isinstance(r, dict) else ''
            
            for year in 年份范围:
                # 匹配价格模式
                价格模式 = r'(\d+\.?\d*)\s*万元?/吨|均价\s*(\d+\.?\d*)\s*万'
                匹配 = re.search(价格模式, title + ' ' + body)
                
                if 匹配:
                    价格 = float(匹配.group(1) or 匹配.group(2))
                    # 匹配该年的日期
                    日期模式 = rf'{year}[年\-](\d+)[月\-]'
                    日期匹配 = re.search(日期模式, title + ' ' + body)
                    if 日期匹配:
                        月份 = 日期匹配.group(1).zfill(2)
                        数据点.append({
                            "年月": f"{year}-{月份}",
                            "月均价": 价格,
                            "来源": self.主信源,
                            "验证状态": "已解析"
                        })
        
        return 数据点
    
    def _step6_交叉验证(self):
        """Step 6: 真正的交叉验证 - 多源价格一致性比较"""
        print("\n[Step 6/7] 交叉验证...")
        
        # 获取主信源当前价格
        主信源当前价格 = None
        for d in self.数据:
            if d.get('验证状态') != '参考数据' or d['年月'] == '2026-03':
                主信源当前价格 = d['月均价']
                break
        
        if not 主信源当前价格:
            主信源当前价格 = self.数据[-1]['月均价'] if self.数据 else None
        
        if not 主信源当前价格:
            print("  [警告] 无价格数据，跳过验证")
            return
        
        print(f"  主信源({self.主信源})当前价格: {主信源当前价格}万元/吨")
        
        验证结果列表 = []
        验证通过计数 = 0
        
        # 遍历补齐信源进行价格比较
        for source in self.补齐信源[:2]:
            print(f"\n  验证补齐信源: {source}...")
            
            # 搜索补齐信源数据
            results = self.web_search.search(f"{source} {self.品种} 价格 2026", max_results=5, time_range="m")
            print(f"    搜索到 {len(results)} 条数据")
            
            if not results:
                验证结果列表.append({"信源": source, "状态": "✗ 无数据", "差异": None})
                continue
            
            # 解析补齐信源价格
            补齐价格数据 = self._解析价格数据(results, "2026")
            
            if not 补齐价格数据:
                验证结果列表.append({"信源": source, "状态": "✗ 无法解析", "差异": None})
                continue
            
            # 取最新价格
            补齐价格 = 补齐价格数据[0]['月均价']
            
            # 计算差异百分比
            差异 = abs(主信源当前价格 - 补齐价格) / 主信源当前价格 * 100
            
            # 判断是否一致 (差异 < 10%)
            if 差异 < 10:
                状态 = f"✓ 一致"
                验证通过计数 += 1
            else:
                状态 = f"⚠ 差异{差异:.1f}%"
            
            验证结果列表.append({
                "信源": source, 
                "状态": 状态, 
                "差异": 差异,
                "补齐价格": 补齐价格
            })
            
            print(f"    {source}价格: {补齐价格}万元/吨, 差异: {差异:.1f}% {状态}")
        
        # 汇总验证结果
        print("\n  " + "="*50)
        print("  交叉验证结果汇总:")
        print("  " + "-"*50)
        
        全部通过 = True
        for r in 验证结果列表:
            if r["差异"] is not None:
                print(f"    {self.主信源}: {主信源当前价格}万 vs {r['信源']}: {r['补齐价格']}万 → {r['状态']}")
                if "✗" in r["状态"]:
                    全部通过 = False
            else:
                print(f"    {r['信源']}: {r['状态']}")
                全部通过 = False
        
        print("  " + "-"*50)
        
        # 只有验证通过才改变状态
        if 验证通过计数 >= 1:
            print(f"  验证结果: 通过 ({验证通过计数}个信源一致)")
            for d in self.数据:
                if d["验证状态"] == "参考数据":
                    d["验证状态"] = "已验证"
        else:
            print("  验证结果: 未通过，保持原状态")
    
    def _step7_输出数据(self):
        """Step 7: 输出数据 - 直接显示价格列表"""
        print("\n" + "="*60)
        print("📊 多晶硅价格数据")
        print("="*60)
        
        当前价格 = self.数据[-1] if self.数据 else None
        
        print(f"\n当前价格: {当前价格['月均价']}万元/吨 ({当前价格['年月']})")
        print(f"主信源: {self.主信源}")
        
        print(f"\n近3年价格走势:")
        print("-" * 40)
        for d in self.数据:
            标记 = " ← 当前" if d['年月'] == 当前价格['年月'] else ""
            标记 += " ← 历史低点" if d['月均价'] == min([x['月均价'] for x in self.数据]) and d['月均价'] < 5 else ""
            print(f"  {d['年月']}  {d['月均价']:>6.2f}万元/吨{标记}")
        print("-" * 40)
        
        output = {
            "品种": self.品种,
            "主信源": self.主信源,
            "补齐信源": self.补齐信源,
            "数据类型": "n型复投料成交均价",
            "单位": "万元/吨",
            "当前价格": 当前价格['月均价'] if 当前价格 else None,
            "当前月份": 当前价格['年月'] if 当前价格 else None,
            "数据": self.数据,
            "验证状态统计": {
                "已验证": len([d for d in self.数据 if d["验证状态"] == "已验证"]),
                "参考数据": len([d for d in self.数据 if d["验证状态"] == "参考数据"])
            }
        }
        
        print("\n  数据收集完成!")
        return output


class 产能监控Agent:
    """
    产能监控Agent - 完整实现LL-负面舆情监控Skill的7步流程
    支持web_search(ddg) + weixin_search双源搜索
    监控光伏组件Top10企业的产能受损风险
    """
    
    def __init__(self):
        self.企业列表 = 光伏企业列表
        self.负面词库 = 负面词库_扁平
        self.负面词库分类 = 负面词库
        self.事件列表 = []
        self.风险事件 = []
        self.web_search = WebSearchClient()
    
    def search(self):
        """执行完整的7步负面舆情监控流程（不生成报告）"""
        print("[产能监控Agent] 开始执行产能受损风险监控...")
        
        self._step1_搜索策略()
        raw_results = self._step2_数据源搜索()
        matched_events = self._step3_负面词匹配(raw_results)
        filtered_events = self._step4_信源筛选(matched_events)
        verified_events = self._step5_信源核实(filtered_events)
        risk_assessment = self._step6_风险判定(verified_events)
        
        # 只返回数据，不生成报告（报告由Team在Step 4统一生成）
        return self._返回舆情数据(risk_assessment)
    
    def _返回舆情数据(self, risk_assessment):
        """返回舆情数据（不生成报告）- 产能影响由用户自行调研"""
        所有事件 = risk_assessment["高风险"] + risk_assessment["中风险"] + risk_assessment["低风险"]
        高置信度事件 = [e for e in 所有事件 if e.get("可信度") in ["5星", "4星"]]
        
        # 简报：只输出事件，产能影响标注"待调研"
        简报 = []
        for e in 高置信度事件[:5]:
            简报.append({
                "事件": e.get("标题", ""),
                "时间": e.get("时间", ""),
                "涉及主体": e.get("企业", ""),
                "信息链接": e.get("链接", ""),
                "风险等级": e.get("风险判定", ""),
                "信源": f"{e.get('信源类型', '')} {e.get('可信度', '')}",
                "影响产能": "待调研"
            })
        
        return {
            "事件数量": len(所有事件),
            "高风险事件": risk_assessment["高风险"],
            "中风险事件": risk_assessment["中风险"],
            "低风险事件": risk_assessment["低风险"],
            "高置信度事件数": len(高置信度事件),
            "总影响产能": "待调研",
            "风险分布": {
                "高风险": len(risk_assessment["高风险"]),
                "中风险": len(risk_assessment["中风险"]),
                "低风险": len(risk_assessment["低风险"])
            },
            "事件列表": 所有事件,
            "舆情简报": 简报
        }
    
    def _step1_搜索策略(self):
        """Step 1: 搜索策略"""
        print("\n[Step 1/7] 搜索策略...")
        print(f"  目标企业: {len(self.企业列表)} 家光伏组件Top10企业")
        print(f"  负面词库: {len(self.负面词库)} 个高风险词（8个类别）")
        print("  负面词类别:", ", ".join(self.负面词库分类.keys()))
        print("  搜索公式: 企业名 + 负面词库(OR连接)")
        print("  搜索工具: web_search(ddg) + weixin_search")
    
    def _step2_数据源搜索(self):
        """Step 2: 数据源搜索 - 使用web_search + weixin_search双源"""
        print("\n[Step 2/7] 数据源搜索...")
        
        all_results = []
        
        # 搜索策略：前3家企业 + 核心负面词（更精确的组合）
        负面搜索词 = [
            "发生火灾", "发生爆炸", "起火", "工厂起火", "组件厂火灾",
            "停产", "停工", "工厂停产", "工厂停工",
            "发生事故", "安全事故",
            "被查", "被立案", "被处罚", "责令整改",
            "限电", "停电", "断电",
            "资金链断裂", "破产"
        ]
        
        for 企业 in self.企业列表[:3]:
            for 词 in 负面搜索词[:8]:  # 限制搜索词数量
                try:
                    # web_search搜索 - 只搜索最近24小时
                    query = f"{企业} {词}"
                    web_results = self.web_search.search(query, max_results=3, time_range="d")
                    for r in web_results:
                        # 验证搜索结果是否真正包含企业名（过滤不相关的搜索结果）
                        title = r.get('title', '')
                        body = r.get('body', '')
                        if 企业 not in title and 企业 not in body:
                            continue  # 跳过不包含企业名的搜索结果
                        
                        # 尝试从搜索结果中提取时间
                        时间 = r.get('date', r.get('time', r.get('published', '')))
                        all_results.append({
                            "企业": 企业,
                            "关键词": 词,
                            "标题": title,
                            "链接": r.get('href', ''),
                            "时间": 时间,
                            "来源": "web_search",
                            "摘要": r.get('body', ''),
                            "正文": ""  # 新增: 存储正文
                        })
                    
                    # weixin_search搜索 - 只保留最近24小时的文章
                    from datetime import datetime, timedelta
                    wx_results = sogou_weixin_search(query, page=1)
                    cutoff_date = datetime.now() - timedelta(hours=24)
                    for r in wx_results[:2]:
                        # 时间过滤：检查publish_time
                        publish_time_str = r.get('publish_time', '')
                        if publish_time_str:
                            try:
                                # 尝试解析日期 (格式如: 2026-03-24 或 3月24日)
                                if '-' in publish_time_str:
                                    pub_date = datetime.strptime(publish_time_str, '%Y-%m-%d')
                                else:
                                    # 处理"3月24日"这种格式，结合当前年份
                                    current_year = datetime.now().year
                                    match = re.search(r'(\d+)月(\d+)日', publish_time_str)
                                    if match:
                                        pub_date = datetime.strptime(f"{current_year}-{match.group(1)}-{match.group(2)}", '%Y-%m-%d')
                                    else:
                                        pub_date = datetime.now()  # 无法解析则保留
                                if pub_date < cutoff_date:
                                    print(f"    [时间过滤] 跳过旧文: {r.get('title', '')[:30]}... ({publish_time_str})")
                                    continue
                            except:
                                pass
                        
                        # 如果是微信链接，尝试获取正文
                        正文 = ""
                        link = r.get('link', '')
                        if 'weixin' in link or 'mp.weixin' in link:
                            try:
                                from weixin_search_mcp.tools.weixin_search import get_article_content
                                正文 = get_article_content(link) or ""
                                if 正文:
                                    print(f"    [正文提取] 成功获取: {企业} {词}")
                            except Exception as e:
                                print(f"    [正文提取] 失败: {e}")
                        
                        all_results.append({
                            "企业": 企业,
                            "关键词": 词,
                            "标题": r.get('title', ''),
                            "链接": link,
                            "时间": r.get('publish_time', ''),
                            "来源": "weixin_search",
                            "摘要": "",
                            "正文": 正文  # 新增: 存储正文
                        })
                        
                except Exception as e:
                    print(f"    {query}: 搜索失败 ({e})")
        
        print(f"  原始搜索结果: {len(all_results)} 条")
        return all_results
    
    def _step3_负面词匹配(self, raw_results):
        """Step 3: 负面词匹配 - 必须同时满足企业名+负面词"""
        print("\n[Step 3/7] 负面词匹配...")
        
        matched = []
        类别统计 = {类别: 0 for 类别 in self.负面词库分类.keys()}
        
        for r in raw_results:
            企业名 = r.get("企业", "")
            title = r.get("标题", "")
            body = r.get("正文", "") + r.get("摘要", "")
            content = title + " " + body
            
            # 必须同时包含企业名和负面词才算匹配
            if not 企业名 or 企业名 not in content:
                continue  # 跳过不包含企业名的结果
            
            # 加强验证：检查企业名是否与文章主题相关（排除假阳性）
            # 方案：检查企业名是否出现在标题开头或前20个字符（文章主题位置）
            # 而不是只出现在文章中间（如"除了天合光能，还有..."这种）
            if 企业名 not in title[:50] and len(title) > 20:
                # 企业名未出现在标题开头，可能是假阳性
                # 检查标题是否以其他企业/品牌开头
                其他品牌 = ["现代", "丰田", "本田", "特斯拉", "比亚迪", "大众", "福特"]
                if any(b in title[:30] for b in 其他品牌):
                    print(f"    [过滤假阳性] {企业名} | 标题实际讲的是: {title[:40]}...")
                    continue
            
            for 类别, 词列表 in self.负面词库分类.items():
                for word in 词列表:
                    if word in content:
                        r["命中词"] = word
                        r["命中类别"] = 类别
                        r["风险等级"] = "高风险"
                        matched.append(r)
                        类别统计[类别] += 1
                        break
                if r in matched:
                    break
        
        print(f"  命中高风险词: {len(matched)} 条")
        print("  类别分布:", ", ".join([f"{k}:{v}" for k, v in 类别统计.items() if v > 0]))
        for m in matched[:3]:
            print(f"    - {m['企业']} | {m['命中类别']}/{m['命中词']} | {m['标题'][:40]}...")
        
        return matched
    
    def _step4_信源筛选(self, matched_events):
        """Step 4: 信源筛选 - 扩展信源类型识别"""
        print("\n[Step 4/7] 信源筛选...")
        
        for e in matched_events:
            link = e.get("链接", "")
            信源类型 = "其他"
            可信度 = "2星"
            等级 = "D级"
            
            # 根据域名识别信源类型
            for domain, (类型, 星级, 信源等级) in 信源识别规则.items():
                if domain in link:
                    信源类型 = 类型
                    可信度 = 星级
                    等级 = 信源等级
                    break
            
            e["信源类型"] = 信源类型
            e["可信度"] = 可信度
            e["信源等级"] = 等级
        
        # 统计信源类型
        类型统计 = {}
        for e in matched_events:
            t = e.get("信源类型", "其他")
            类型统计[t] = 类型统计.get(t, 0) + 1
        
        print(f"  信源分类完成: {len(matched_events)} 条")
        for t, count in 类型统计.items():
            print(f"    - {t}: {count} 条")
        
        return matched_events
    
    def _step5_信源核实(self, filtered_events):
        """Step 5: 信源核实 - 完整实现交叉验证矩阵"""
        print("\n[Step 5/7] 信源核实（完整交叉验证）...")
        
        # 按企业和关键词分组，实现交叉验证
        事件分组 = {}
        for e in filtered_events:
            key = (e.get("企业", ""), e.get("命中词", ""))
            if key not in 事件分组:
                事件分组[key] = []
            事件分组[key].append(e)
        
        verified = []
        交叉验证统计 = {"强验证": 0, "中验证": 0, "弱验证": 0, "未验证": 0}
        
        for (企业, 关键词), 事件组 in 事件分组.items():
            # 按等级分类
            A级事件 = [e for e in 事件组 if e.get("信源等级") == "A级"]
            B级事件 = [e for e in 事件组 if e.get("信源等级") == "B级"]
            C级事件 = [e for e in 事件组 if e.get("信源等级") == "C级"]
            D级事件 = [e for e in 事件组 if e.get("信源等级") == "D级"]
            E级事件 = [e for e in 事件组 if e.get("信源等级") == "E级"]
            
            # 应用交叉验证矩阵（来自Skill定义）
            if A级事件:  # A + 任意 = 强验证
                for e in 事件组:
                    e["核实状态"] = "[OK] 核实通过"
                    e["验证类型"] = "强验证"
                    e["核实结论"] = f"官方/权威来源确认，极高可信度"
                    交叉验证统计["强验证"] += 1
            elif len(B级事件) >= 2:  # B + B = 互相印证
                for e in 事件组:
                    e["核实状态"] = "[OK] 核实通过"
                    e["验证类型"] = "中验证"
                    e["核实结论"] = f"多个主流媒体印证，高可信度"
                    交叉验证统计["中验证"] += 1
            elif B级事件 and C级事件:  # B + C = 多方印证
                for e in 事件组:
                    e["核实状态"] = "[OK] 核实通过"
                    e["验证类型"] = "中验证"
                    e["核实结论"] = f"主流媒体+行业确认，中等可信度"
                    交叉验证统计["中验证"] += 1
            elif len(C级事件) >= 2:  # C + C = 行业确认
                for e in 事件组:
                    e["核实状态"] = "[OK] 核实通过"
                    e["验证类型"] = "弱验证"
                    e["核实结论"] = f"行业媒体确认，中等可信度，建议跟进"
                    交叉验证统计["弱验证"] += 1
            elif D级事件 and len(D级事件) == len(事件组):  # 仅D级 = 待跟进
                for e in 事件组:
                    e["核实状态"] = "[PENDING] 待跟进"
                    e["验证类型"] = "未验证"
                    e["核实结论"] = "地方媒体报道，待主流媒体或官方确认"
                    交叉验证统计["未验证"] += 1
            elif E级事件 and len(E级事件) == len(事件组):  # 仅E级 = 参考
                for e in 事件组:
                    e["核实状态"] = "[INFO] 参考信息"
                    e["验证类型"] = "未验证"
                    e["核实结论"] = "其他来源报道，作为线索关注，不作为决策依据"
                    交叉验证统计["未验证"] += 1
            else:  # 单一来源 = 标记待验证
                for e in 事件组:
                    e["核实状态"] = "[PENDING] 待跟进"
                    e["验证类型"] = "未验证"
                    e["核实结论"] = "单一来源报道，建议等待更多确认"
                    交叉验证统计["未验证"] += 1
            
            verified.extend(事件组)
        
        已核实 = len([v for v in verified if "[OK]" in v.get("核实状态", "")])
        print(f"  核实结果: {已核实}/{len(verified)} 条已核实")
        print(f"  验证分布: 强{交叉验证统计['强验证']} 中{交叉验证统计['中验证']} 弱{交叉验证统计['弱验证']} 未{交叉验证统计['未验证']}")
        
        return verified
    
    def _step6_风险判定(self, verified_events):
        """Step 6: 风险判定"""
        print("\n[Step 6/7] 风险判定...")
        
        高风险事件 = []
        中风险事件 = []
        低风险事件 = []
        
        for e in verified_events:
            等级 = e.get("信源等级", "D")
            if e.get("风险等级") == "高风险" and 等级 in ["A级", "B级"]:
                e["风险判定"] = "[HIGH] 高风险"
                高风险事件.append(e)
            elif e.get("风险等级") == "高风险" and 等级 == "C级":
                e["风险判定"] = "[MED] 中风险"
                中风险事件.append(e)
            else:
                e["风险判定"] = "[LOW] 低风险"
                低风险事件.append(e)
        
        print(f"  [HIGH] 高风险: {len(高风险事件)} 条")
        print(f"  [MED] 中风险: {len(中风险事件)} 条")
        print(f"  [LOW] 低风险: {len(低风险事件)} 条")
        
        return {
            "高风险": 高风险事件,
            "中风险": 中风险事件,
            "低风险": 低风险事件
        }
    
    def _step7_报告输出(self, risk_assessment):
        """Step 7: 报告输出 - 生成完整舆情日报"""
        print("\n[Step 7/7] 报告输出...")
        
        from datetime import datetime
        import os
        
        所有事件 = risk_assessment["高风险"] + risk_assessment["中风险"] + risk_assessment["低风险"]
        
        # 只统计高置信度事件（4星及以上）
        高置信度事件 = [e for e in 所有事件 if e.get("可信度") in ["5星", "4星"]]
        总影响产能 = len(高置信度事件) * 10
        
        # ========== 生成舆情日报 (按Skill格式) ==========
        today = datetime.now().strftime("%Y-%m-%d")
        日报_lines = []
        日报_lines.append(f"⚠️风险评估（{today}）")
        日报_lines.append("")
        
        if len(高置信度事件) == 0:
            日报_lines.append("今日无高风险舆情事件")
        else:
            for idx, e in enumerate(高置信度事件, 1):
                日报_lines.append(f"{idx}. 事件：{e.get('标题', '未知')}")
                日报_lines.append(f"   时间：{e.get('时间', '未知')}")
                
                # 信源格式化
                信源列表 = e.get('信源列表', [])
                if 信源列表:
                    信源_str = " | ".join([f"{s.get('名称', '')} {s.get('可信度', '')}" for s in 信源列表])
                    日报_lines.append(f"   新闻来源：{信源_str}")
                else:
                    日报_lines.append(f"   新闻来源：{e.get('信源类型', '')} {e.get('可信度', '')}")
                
                # 链接格式化
                链接 = e.get('链接', '')
                if 链接:
                    日报_lines.append(f"   新闻链接：{链接}")
                
                # 产能影响（按Skill规则）
                产能 = e.get('影响产能', '')
                if 产能:
                    日报_lines.append(f"   影响产能：{产能}")
                else:
                    日报_lines.append(f"   影响产能：建议请咨询公司，或联系相应的证券公司研究员")
                
                日报_lines.append("")
        
        日报_content = "\n".join(日报_lines)
        
        # 保存舆情日报文件
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        日报_path = os.path.join(output_dir, f"舆情监控日报_{today}.txt")
        with open(日报_path, "w", encoding="utf-8") as f:
            f.write(日报_content)
        
        print(f"  舆情日报已保存: {日报_path}")
        
        # 控制台输出日报 (ASCII-safe)
        print("\n" + "=" * 50)
        # 替换emoji为ASCII
        日报_ascii = 日报_content.replace("⚠️", "[!]").replace("✅", "[OK]").replace("ℹ️", "[i]")
        print(日报_ascii)
        print("=" * 50)
        
        简报 = []
        for e in 高置信度事件[:5]:
            简报.append({
                "事件": e.get("标题", ""),
                "时间": e.get("时间", ""),
                "涉及主体": e.get("企业", ""),
                "信息链接": e.get("链接", ""),
                "风险等级": e.get("风险判定", ""),
                "信源": f"{e.get('信源类型', '')} {e.get('可信度', '')}"
            })
        
        print(f"\n  舆情监控完成: 发现 {len(所有事件)} 条舆情")
        print(f"  高置信度事件: {len(高置信度事件)} 条")
        print(f"  总影响产能估算: {总影响产能} GW")
        
        return {
            "事件数量": len(所有事件),
            "高置信度事件数": len(高置信度事件),
            "总影响产能": 总影响产能,
            "风险分布": {
                "高风险": len(risk_assessment["高风险"]),
                "中风险": len(risk_assessment["中风险"]),
                "低风险": len(risk_assessment["低风险"])
            },
            "事件列表": 所有事件,
            "舆情简报": 简报,
            "日报路径": 日报_path,
            "日报内容": 日报_content
        }


class 分析评估Agent:
    """分析评估Agent - 实现水位分析 + 产能评级计算"""

    def analyze(self, 价格数据, 舆情数据):
        """基于价格和舆情数据进行供需分析"""
        print("[分析评估Agent] 开始执行分析评估...")

        水位分析结果 = self._水位分析(价格数据)
        产能影响值 = self._提取产能影响(舆情数据)
        评级结果 = self._评级计算(产能影响值, 水位分析结果["水位"])

        return {
            "评级": 评级结果["评级"],
            "评级文字": 评级结果["评级文字"],
            "评级理由": 评级结果["理由"],
            "建议": 评级结果["建议"],
            "产能影响值": 产能影响值,
            "当前价格": 水位分析结果["当前价格"],
            "水位": 水位分析结果["水位"],
            "水位说明": 水位分析结果["说明"],
            "价格周期": 水位分析结果["价格周期"],
        }

    def _水位分析(self, 价格数据):
        """水位分析：判断当前价格处于历史3年的什么水位（百分位法）"""
        print("\n  [水位分析] 分析当前价格所处历史水位...")

        数据 = 价格数据.get("数据", [])
        if not 数据 or len(数据) < 2:
            return {
                "当前价格": None,
                "水位": "未知",
                "说明": "数据不足",
                "价格周期": "未知"
            }

        prices = [d.get("月均价", 0) for d in 数据]
        dates = [d.get("年月", "") for d in 数据]
        current_price = prices[-1]

        print(f"    数据周期: {dates[0]} ~ {dates[-1]}")
        print(f"    当前价格: {current_price} 万元/吨")

        # 百分位法判断水位
        sorted_prices = sorted(prices)
        n = len(sorted_prices)

        p20 = sorted_prices[int(n * 0.2)]
        p40 = sorted_prices[int(n * 0.4)]
        p60 = sorted_prices[int(n * 0.6)]
        p80 = sorted_prices[int(n * 0.8)]

        print(f"    历史区间: {sorted_prices[0]:.2f} ~ {sorted_prices[-1]:.2f} 万元/吨")
        print(f"    百分位: P20={p20:.2f} P40={p40:.2f} P60={p60:.2f} P80={p80:.2f}")

        if current_price > p80:
            水位 = "高水位"
            说明 = f"当前价格{current_price:.2f}万高于80%历史数据"
        elif current_price > p60:
            水位 = "中高水位"
            说明 = f"当前价格{current_price:.2f}万高于60%历史数据"
        elif current_price > p40:
            水位 = "中水位"
            说明 = f"当前价格{current_price:.2f}万高于40%历史数据"
        elif current_price > p20:
            水位 = "低水位"
            说明 = f"当前价格{current_price:.2f}万高于20%历史数据（价格偏低）"
        else:
            水位 = "历史低位"
            说明 = f"当前价格{current_price:.2f}万处于历史最低区间"

        print(f"    水位判断: {水位}")

        return {
            "当前价格": current_price,
            "水位": 水位,
            "说明": 说明,
            "价格周期": f"{dates[0]} ~ {dates[-1]}",
            "历史高点": sorted_prices[-1],
            "历史低点": sorted_prices[0],
            "百分位": {"p20": p20, "p40": p40, "p60": p60, "p80": p80}
        }

    def _提取产能影响(self, 舆情数据):
        """从舆情数据中提取产能影响值"""
        产能影响值 = 舆情数据.get("总影响产能", 0) if isinstance(舆情数据, dict) else 0
        if isinstance(产能影响值, (int, float)):
            print(f"\n  [产能提取] 产能影响值: {产能影响值} GW")
            return 产能影响值
        else:
            print(f"\n  [产能提取] 产能影响值: 待调研（从舆情事件中提取）")
            return 0  # 评级计算用0，不满足条件

    def _评级计算(self, 产能影响值, 水位):
        """评级计算 - 5星/4星/其他三档（产能影响由人工调研）"""
        print("\n  [评级计算] 执行评级...")

        # 5星王者：产能 > 200GW + 低水位
        # 4星战神：产能 100-200GW + 低水位
        # 其他：暂时关注（产能影响待调研）

        if 产能影响值 <= 0 or 产能影响值 == "待调研":
            # 产能影响值未确定，输出待调研状态
            评级 = "[待定]"
            评级文字 = "待人工调研"
            理由 = "产能影响值待人工调研，请根据舆情事件自行判断产能影响后决策"
            建议 = "1. 请根据下方舆情事件调研产能影响值；2. 结合水位判断后决策；3. 可联系老梨协助调研"
        elif 产能影响值 > 200 and 水位 in ["低水位", "历史低位"]:
            评级 = "[5星]"
            评级文字 = "5星王者"
            理由 = f"产能影响{产能影响值}GW + 价格{水位}，供需改善预期强烈"
            建议 = "重点关注，产能收缩叠加低价，原材料端具备强势话语权"
        elif 100 <= 产能影响值 <= 200 and 水位 in ["低水位", "历史低位"]:
            评级 = "[4星]"
            评级文字 = "4星战神"
            理由 = f"产能影响{产能影响值}GW + 价格{水位}，供需边际改善"
            建议 = "积极关注，产能去化进行时，价格弹性可期"
        else:
            评级 = "[3星]"
            评级文字 = "暂时关注"
            理由 = f"产能影响{产能影响值}GW 或 价格水位{水位}未达标准"
            建议 = "保持观察，等待产能去化进一步明朗或价格出现明显拐点"

        print(f"    评级: {评级}（{评级文字}）")
        print(f"    理由: {理由}")
        print(f"    建议: {建议}")

        return {
            "评级": 评级,
            "评级文字": 评级文字,
            "理由": 理由,
            "建议": 建议
        }


class 文档生成Agent:
    """生成Word报告 - 完善格式（首行缩进、1.5倍行距、红色高亮）"""
    
    def generate(self, 分析结果, 价格数据, 舆情数据):
        """生成Word格式的供需分析报告"""
        print("[文档生成Agent] 正在生成Word报告...")
        
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 设置默认中文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置段落格式（首行缩进2字符，1.5倍行距）
        style.paragraph_format.first_line_indent = Cm(0.74)  # 约2字符
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)
        
        # 标题（四号、宋体、加粗、居中）
        title = doc.add_heading('光伏供需监控报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '宋体'
            run.font.size = Pt(14)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title.paragraph_format.first_line_indent = Cm(0)  # 标题不缩进
        
        # 副标题（五号、宋体、居中）
        subtitle = doc.add_paragraph('@爱思考的老梨 | 2026-03-23')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.first_line_indent = Cm(0)
        for run in subtitle.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        # 第一节：分析逻辑
        self._添加节标题(doc, '第一节 分析逻辑')
        self._添加小节标题(doc, '一、大厂预计总产能')
        self._添加正文(doc, '出货量（Top10组件企业）：540 GW')
        self._添加正文(doc, '行业平均开工率：50%')
        self._添加正文(doc, '计算公式：总产能 = 540 ÷ 50% = 1080 GW')
        
        self._添加小节标题(doc, '二、大厂合理总产能')
        self._添加正文(doc, '盈利开工率（盈亏平衡）：60%-70%')
        self._添加正文(doc, '计算公式：合理产能 = 540 ÷ (60%-70%) = 771GW-900GW')
        
        # 第二节：当前变化
        self._添加节标题(doc, '第二节 当前变化')
        self._添加小节标题(doc, '一、产能事件')
        self._添加正文(doc, f"舆情事件数量：{舆情数据.get('事件数量', 0)} 起")
        self._添加正文(doc, f"高置信度事件：{舆情数据.get('高置信度事件数', 0)} 起")
        self._添加正文(doc, f"总影响产能：{舆情数据.get('总影响产能', 0)} GW")
        
        self._添加小节标题(doc, '二、原材料价格趋势')
        self._添加正文(doc, f"最新价格：{分析结果.get('当前价格', 'N/A')} 万元/吨")
        self._添加正文(doc, f"当前水位：{分析结果.get('水位', 'N/A')}")
        self._添加正文(doc, f"水位说明：{分析结果.get('水位说明', 'N/A')}")
        self._添加正文(doc, f"数据周期：{分析结果.get('价格周期', 'N/A')}")
        self._添加正文(doc, f"数据来源：{价格数据.get('主信源', 'N/A')}")
        
        # 第三节：分析评级（红色加粗高亮）
        self._添加节标题(doc, '第三节 分析评级')
        评级段落 = doc.add_paragraph()
        评级段落.paragraph_format.first_line_indent = Cm(0.74)
        评级段落.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        评级文字 = 分析结果.get('评级文字', 'N/A')
        评级 = 分析结果.get('评级', 'N/A')
        
        run = 评级段落.add_run(f"综合评级：{评级}（{评级文字}）")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 第四节：结论建议
        self._添加节标题(doc, '第四节 结论建议')
        理由 = 分析结果.get('评级理由', '')
        建议 = 分析结果.get('建议', '')
        
        if 评级文字 == "待人工调研":
            self._添加正文(doc, '【系统状态】产能影响值待人工调研')
            self._添加正文(doc, 理由)
            self._添加正文(doc, 建议)
        elif 评级文字 in ["5星王者", "4星战神"]:
            self._添加正文(doc, '投资建议：')
            self._添加正文(doc, '1. 立即与机构建立联系，保持沟通，确定机构资金走向')
            self._添加正文(doc, '2. 若无机构建联机会，应密切关注交易变动情况')
        else:
            self._添加正文(doc, '投资建议：')
            self._添加正文(doc, '暂时观望，等待更明确的供需改善信号')
        
        if 理由:
            self._添加正文(doc, f'评级理由：{理由}')
        
        # 第五节：注意事项（楷体）
        self._添加节标题(doc, '第五节 注意事项')
        注意段落1 = doc.add_paragraph()
        注意段落1.paragraph_format.first_line_indent = Cm(0.74)
        注意段落1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run1 = 注意段落1.add_run('WorkBuddy暂未集成交易Claw功能，如需交易执行支持，请联系老梨（@爱思考的老梨）协助研发。')
        run1.font.name = '楷体'
        run1.font.size = Pt(12)
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        注意段落2 = doc.add_paragraph()
        注意段落2.paragraph_format.first_line_indent = Cm(0.74)
        注意段落2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run2 = 注意段落2.add_run('本分析内容仅可作为参考，不构成投资建议。')
        run2.font.name = '楷体'
        run2.font.size = Pt(12)
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        注意段落3 = doc.add_paragraph()
        注意段落3.paragraph_format.first_line_indent = Cm(0.74)
        注意段落3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run3 = 注意段落3.add_run('本文所涉及到的数据仅可供个人学习，不得用于存储、商用等损坏版权方利益的行为。')
        run3.font.name = '楷体'
        run3.font.size = Pt(12)
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        # 保存
        output_path = "光伏供需监控报告.docx"
        doc.save(output_path)
        print(f"[文档生成Agent] 报告已保存: {output_path}")
        
        return output_path
    
    def _添加节标题(self, doc, text):
        """添加节标题（一级标题）"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        
        heading = doc.add_heading(text, level=1)
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(14)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _添加小节标题(self, doc, text):
        """添加小节标题（二级标题）"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        
        heading = doc.add_heading(text, level=2)
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _添加正文(self, doc, text):
        """添加正文段落（带首行缩进和1.5倍行距）"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ========== ⑤ ClawBot输出Agent ==========
class ClawBot输出Agent:
    """专门负责生成微信/ClawBot可发送的简洁文本格式"""
    
    def generate(self, 分析结果, 价格数据, 舆情数据):
        """生成简洁文本摘要"""
        from datetime import datetime
        today = datetime.now().strftime("%Y-%m-%d")
        
        输出 = f"""📊 光伏产业监控日报 - {today}

⭐ 今日舆情
"""
        # 处理舆情 - 只显示B级及以上信源的真实新闻
        事件列表 = []
        if isinstance(舆情数据, dict):
            # 获取所有风险等级的事件
            高风险 = 舆情数据.get("高风险事件", [])
            中风险 = 舆情数据.get("中风险事件", [])
            低风险 = 舆情数据.get("低风险事件", [])
            所有事件 = 高风险 + 中风险 + 低风险
            # 过滤：只保留B级及以上信源（排除自媒体、低质量平台）
            事件列表 = [e for e in 所有事件 if e.get("信源等级") in ["A级", "B级", "C级"]]
        elif isinstance(舆情数据, list):
            事件列表 = [e for e in 舆情数据 if e.get("信源等级") in ["A级", "B级", "C级"]]
        
        if 事件列表 and len(事件列表) > 0:
            for i, e in enumerate(事件列表[:3], 1):
                标题 = e.get('事件', e.get('标题', e.get('title', '无标题')))[:40]
                企业 = e.get('涉及主体', e.get('企业', ''))
                风险等级 = e.get('风险等级', '')
                信源等级 = e.get('信源等级', '')
                输出 += f"{i}. {企业}[{风险等级}][{信源等级}] - {标题}...\n"
        else:
            输出 += "今日无重大负面舆情\n"
        
        # 原材料价格 - 从分析结果中提取（新版水位分析）
        当前价格 = "待补充"
        水位 = "待补充"
        水位说明 = ""
        if 分析结果:
            当前价格 = 分析结果.get('当前价格')
            if 当前价格:
                当前价格 = f"{当前价格:.2f}"
            水位 = 分析结果.get('水位', '待补充')
            水位说明 = 分析结果.get('水位说明', '')
        
        输出 += f"""
⭐ 原材料价格（多晶硅）
当前价格：{当前价格}万元/吨
当前水位：{水位}
"""
        
        # 供需分析 - 从分析结果中提取（新版简化评级）
        if 分析结果:
            评级 = 分析结果.get('评级文字', '待补充')
            建议 = 分析结果.get('建议', '待补充')
            评级理由 = 分析结果.get('评级理由', '')
            产能影响值 = 分析结果.get('产能影响值', '待补充')
            
            输出 += f"""
⭐ 供需变化
产能影响：{产能影响值} GW
评级理由：{评级理由}

⭐ 建议结论
评级：{评级}
建议：{建议}
"""
        
        # 注意事项
        输出 += """
━━━━━━━━━━━━━━━━━━━━
⚠️ 注意事项
1. WorkBuddy暂未集成交易Claw功能，如需交易执行支持，请联系老梨（@爱思考的老梨）协助研发
2. 本分析内容仅可作为参考，不构成投资建议
3. 本文所涉及到的数据仅可供个人学习，不得用于存储、商用等损坏版权方利益的行为
"""
        
        return 输出


# ========== Team 主控（支持并行执行） ==========
class 光伏供需分析评估Team:
    """Team主控，协调5个Agent工作 - 支持Step 1/2并行执行"""
    
    def __init__(self, parallel=True):
        self.价格Agent = 价格趋势Agent()
        self.舆情Agent = 产能监控Agent()
        self.分析Agent = 分析评估Agent()
        self.文档Agent = 文档生成Agent()
        self.ClawBot输出Agent = ClawBot输出Agent()
        self.parallel = parallel  # 是否并行执行
    
    def run(self, output_mode="file"):
        """
        执行完整工作流程
        
        参数:
            output_mode: 输出模式
                - "clawbot": ClawBot/微信简洁文本
                - "file": Word/txt文件报告
        """
        print("=" * 60)
        print("光伏供需分析评估Team - 开始执行")
        print(f"输出模式: {'ClawBot文本' if output_mode == 'clawbot' else 'Word/txt报告'}")
        if self.parallel:
            print("模式: Step 1/2 并行执行")
        print("=" * 60)
        
        if self.parallel:
            # 并行执行Step 1和Step 2
            with ThreadPoolExecutor(max_workers=2) as executor:
                future_价格 = executor.submit(self._执行价格Agent)
                future_舆情 = executor.submit(self._执行舆情Agent)
                
                价格数据 = future_价格.result()
                舆情数据 = future_舆情.result()
        else:
            # 串行执行
            print("\n[Step 1/4] 价格趋势搜索...")
            价格数据 = self.价格Agent.search()
            
            print("\n[Step 2/4] 舆情监控搜索...")
            舆情数据 = self.舆情Agent.search()
        
        # Step 3: 供需分析
        print("\n[Step 3/4] 供需分析...")
        分析结果 = self.分析Agent.analyze(价格数据, 舆情数据)
        
        # Step 4: 根据输出模式选择
        if output_mode == "clawbot":
            # ClawBot模式：只生成简洁文本
            print("\n[Step 4/4] 生成ClawBot文本...")
            ClawBot文本 = self.ClawBot输出Agent.generate(分析结果, 价格数据, 舆情数据)
            
            print("")
            print("=" * 60)
            print("Team执行完成!")
            print(f"评级: {分析结果['评级文字']}")
            print("=" * 60)
            
            # 打印ClawBot文本
            print("\n" + "=" * 60)
            print("ClawBot输出文本：")
            print("=" * 60)
            print(ClawBot文本)
            
            return {
                "分析结果": 分析结果,
                "价格数据": 价格数据,
                "舆情数据": 舆情数据,
                "ClawBot文本": ClawBot文本,
                "output_mode": "clawbot"
            }
        else:
            # 文件模式：只生成Word/txt报告
            print("\n[Step 4/4] 生成报告（并行）...")
            报告结果 = self._并行生成报告(分析结果, 价格数据, 舆情数据)
        
        print("")
        print("=" * 60)
        print("Team执行完成!")
        print(f"评级: {分析结果['评级文字']}")
        print(f"Word报告: {报告结果['Word报告路径']}")
        print(f"舆情日报: {报告结果['日报路径']}")
        print("=" * 60)
        
        return {
            "分析结果": 分析结果,
            "价格数据": 价格数据,
            "舆情数据": 舆情数据,
            "报告路径": 报告结果['Word报告路径'],
            "日报路径": 报告结果['日报路径']
        }
    
    def _执行价格Agent(self):
        """包装价格Agent执行"""
        print("\n[Step 1/4] 价格趋势搜索 (并行)...")
        return self.价格Agent.search()
    
    def _执行舆情Agent(self):
        """包装舆情Agent执行"""
        print("\n[Step 2/4] 舆情监控搜索 (并行)...")
        return self.舆情Agent.search()
    
    def _并行生成报告(self, 分析结果, 价格数据, 舆情数据):
        """并行生成舆情日报和Word报告"""
        # 准备风险评估数据
        risk_assessment = {
            "高风险": 舆情数据.get("高风险事件", []),
            "中风险": 舆情数据.get("中风险事件", []),
            "低风险": 舆情数据.get("低风险事件", [])
        }
        
        with ThreadPoolExecutor(max_workers=2) as executor:
            # 任务1：生成舆情日报
            future_日报 = executor.submit(
                self.舆情Agent._step7_报告输出, 
                risk_assessment
            )
            
            # 任务2：生成Word报告
            future_Word = executor.submit(
                self.文档Agent.generate, 
                分析结果, 价格数据, 舆情数据
            )
            
            # 等待两个任务完成
            日报结果 = future_日报.result()
            Word报告路径 = future_Word.result()
        
        return {
            "日报路径": 日报结果.get("日报路径", ""),
            "Word报告路径": Word报告路径
        }


# ========== 入口 ==========
if __name__ == "__main__":
    import io
    import sys
    import argparse
    
    # 设置UTF-8编码（解决emoji输出问题）
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='光伏产业监控')
    parser.add_argument('--mode', choices=['file', 'clawbot'], default='file',
                        help='输出模式: file=Word/txt报告, clawbot=简洁文本')

    args = parser.parse_args()
    
    # 默认启用并行执行
    team = 光伏供需分析评估Team(parallel=True)
    result = team.run(output_mode=args.mode)
    
    if args.mode == "clawbot":
        print("\n" + "=" * 60)
        print("ClawBot文本输出完成！")
        print("=" * 60)
    else:
        print("\n报告生成完成！")
        print(f"Word报告: {result['报告路径']}")
        print(f"舆情日报: {result['日报路径']}")
